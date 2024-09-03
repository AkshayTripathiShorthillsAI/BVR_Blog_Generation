import os
from groq import Groq
from dotenv import load_dotenv
from docx import Document
import re  # Import the re module

class ReviewSnippets:
    def __init__(self):
        self.df_pid = None
        self.features_df = None
        load_dotenv()

    # Function to get the blog content
    def get_blog(self, topic):
        client = Groq(
            api_key=os.getenv("GROQ_API_KEY"),
        )

        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": f'''Generate blog for my e-commerce site BestViewsReviews on .It is a GenAI platform that analyses unstructured e-commerce product information like customer reviews, blogs, manufacturer written product description, etc. and generates insights and product recommendations.
                    Following will be the structure of the blog:
                    1. **Introduction**: Brief overview of Amazon Prime Day and Importance and popularity
                    2. **Key Dates and Times**: Official dates and times for Prime Day 2024, Time zones and global variations, Duration of the event
                    3. **How to Prepare**: Sign up for Amazon Prime membership, Download the Amazon app, Enable notifications for deals,Create a wishlist, Set a budget.
                    4. **How to Save More Money on Prime Day**: Utilizing Amazon's Lightning Deals and Daily Deals, Cashback and rewards programs(Amazon credit cards, Partnered bank offers),Combining discounts with Amazon coupons.
                    5. **Top Deals to Look Out For**: Predicted best deals based on previous years, Specific product categories to watch, Brands known for offering significant discounts.
                    6. **Insider Tips for a Successful Prime Day**: How to spot the best deals, Using Amazon Assistant and other tools, Comparing prices and using price trackers, Taking advantage of Amazon’s return policy, Leveraging credit card rewards and cashback offers.
                    7. **Exclusive Benefits for Prime Members**:Free shipping,Prime Video, Music, and Reading offers,Amazon Fresh and Whole Foods deals,Early access to other events and sales
                    8. ** Insider Tips for a Successful Prime Day**: Maximizing savings with Amazon Warehouse and Renewed deals, Comparing prices with other retailers, Using price trackers and comparison tools,Timing your purchases for the best deals.
                    9. **Conclusion**: Recap of key points, Encouragement to make the most of the event.
                    topic: {topic}'''
                },
                {
                    "role": "user",
                    "content": f'''Topic: {topic}'''
                },
            ],
            model="llama3-70b-8192",
        )

        return chat_completion.choices[0].message.content

    # Function to save content to a .doc file with bold formatting
    def save_to_doc(self, content, filename):
        doc = Document()
        paragraph = doc.add_paragraph()
        
        # Split the content by double asterisks to identify bold text
        parts = re.split(r'(\*\*.*?\*\*)', content)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Add bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Add normal text
                paragraph.add_run(part)
        
        doc.save(filename)

if __name__ == "__main__":
    extractor = ReviewSnippets()
    topic = "Amazon Prime Big Deal Days Sale"
    blog_content = extractor.get_blog(topic)
    print(blog_content)

    # Save the blog content to a .doc file with bold formatting
    output_dir = "output"
    os.mkdir(output_dir, exist_ok=True)

    filename = os.path.join(output_dir, "Blog_Content.docx")
    extractor.save_to_doc(blog_content, filename)
    print(f"Blog content saved to {filename}")
