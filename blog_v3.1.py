import os
from groq import Groq
from dotenv import load_dotenv
from docx import Document
import re  # Import the re module

class ReviewSnippets:
    def __init__(self):
        load_dotenv()
        self.client = Groq(api_key=os.getenv("GROQ_API_KEY"))

    def get_blog(self, topic, structure,description):
        chat_completion = self.client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": f'''Generate a blog for my e-commerce site BestViewsReviews for Amazon Products. It is a GenAI platform that analyzes unstructured e-commerce product information like customer reviews, blogs, manufacturer written product descriptions, etc., and generates insights and product recommendations. 
                    Following is the structure of the blog:
                    {structure}
                    Topic: {topic}
product description: {description}
'''
                },
                {
                    "role": "user",
                    "content": f'Topic: {topic}'
                },
            ],
            model="llama3-70b-8192",
        )

        return chat_completion.choices[0].message.content

    def save_to_doc(self, content, filename):
        doc = Document()
        paragraph = doc.add_paragraph()
        
        # Split the content by double asterisks to identify bold text
        parts = re.split(r'(\*\*.*?\*\*)', content)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Add bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Add normal text
                paragraph.add_run(part)
        
        doc.save(filename)

if __name__ == "__main__":
    extractor = ReviewSnippets()
    topic = "Best Food Processors"
    structure = '''Introduction

    Top picks based on aspects – two products each aspect (include at least 4 to 5 aspects ).Give each product description in atlead 2 paragraphs.

    Things to keep in mind while choosing a food processor (6 to 7 factors to consider while buying a food processor and explain them in detail atleat in 2 paragraphs)
    Explain each FAQ and each "Things to keep in mind" in detail each.Explain top picks product in atleast 2 paragraphs.
    Expert tips on how to maintain a food processor

    FAQs 
    - Can a food processor replace a blender?
    - Can I use a food processor for hot ingredients?
    - What are 3 things you can do with a food processor?
    - Which is better: food processor or juicer?
    - What size of food processor do I need?

    Has context menu
    '''


    description = '''[
    "MORE FUNCTIONALITY: The Ninja Professional Plus Kitchen System with Auto-iQ features a new modern design and more functionality than Ninja's original Kitchen System. (Versus BL770 based on the number of available blending programs.)",
    "POWERFUL CRUSHING: Ninja Total Crushing Blades give you perfectly crushed ice for your smoothies and frozen drinks with 1400 peak watts of professional power.",
    "FOOD PROCESSING: The 8-cup Precision Processor Bowl provides precision processing for even chopping and smooth purees.",
    "5 VERSATILE FUNCTIONS: 5 preset Auto-iQ programs allow you to create smoothies, frozen drinks, nutrient extractions, chopped mixtures, and dough, all at the touch of a button. Extract a drink containing vitamins and nutrients from fruits and vegetables.",
    "AUTO-IQ TECHNOLOGY: take the guesswork out of drink making with intelligent programs that combine unique timed pulsing, blending, and pausing patterns that do the work for you.",
    "XL CAPACITY: The 72 oz Total Crushing Pitcher is great for making large batches for the whole family. (64 oz max liquid capacity).",
    "ON-THE-GO CONVENIENCE: Two 24-oz. single-serve cups with spout lids make it easy to take delicious nutrient-rich smoothies on the go."
  ]'''



    blog_content = extractor.get_blog(topic, structure=structure, description=description) 
    print(blog_content)

    # Save the blog content to a .doc file with bold formatting
    output_dir = "output"
    os.mkdir(output_dir, exist_ok=True)

    filename = os.path.join(output_dir, "Blog_Content.docx")
    extractor.save_to_doc(blog_content, filename)
    print(f"Blog content saved to {filename}")
