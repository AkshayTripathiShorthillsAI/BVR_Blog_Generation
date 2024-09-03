import os
from groq import Groq
from dotenv import load_dotenv
from docx import Document
import re  # Import the re module

class ReviewSnippets:
    def __init__(self):
        self.df_pid = None
        self.features_df = None
        load_dotenv()

    # Function to get the blog content
    def get_blog(self, topic):
        client = Groq(
            api_key=os.getenv("GROQ_API_KEY"),
        )

        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": f'''Generate blog for my e-commerce site BestViewsReviews on .It is a GenAI platform that analyses unstructured e-commerce product information like customer reviews, blogs, manufacturer written product description, etc. and generates insights and product recommendations.
                    Following will be the structure of the blog:
                    1. **Introduction**
                    2. **Key Dates and Times**
                    3. **How to Prepare**
                    4. **How to Save More Money on Prime Day**
                    5. **Top Deals to Look Out For**
                    6. **Insider Tips for a Successful Prime Day**
                    7. **Exclusive Benefits for Prime Members**
                    8. ** Insider Tips for a Successful Prime Day**
                    9. **Conclusion**
                    topic: {topic}'''
                },
                {
                    "role": "user",
                    "content": f'''Topic: {topic}'''
                },
            ],
            model="llama3-70b-8192",
        )

        return chat_completion.choices[0].message.content

    # Function to save content to a .doc file with bold formatting
    def save_to_doc(self, content, filename):
        doc = Document()
        paragraph = doc.add_paragraph()
        
        # Split the content by double asterisks to identify bold text
        parts = re.split(r'(\*\*.*?\*\*)', content)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Add bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Add normal text
                paragraph.add_run(part)
        
        doc.save(filename)

if __name__ == "__main__":
    extractor = ReviewSnippets()
    topic = "Amazon Prime Big Deal Days Sale"
    blog_content = extractor.get_blog(topic)
    print(blog_content)

     # Save the blog content to a .doc file with bold formatting
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    filename = os.path.join(output_dir, "Blog_Content.docx")
    extractor.save_to_doc(blog_content, filename)
    print(f"Blog content saved to {filename}")
