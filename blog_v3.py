import os
from groq import Groq
from dotenv import load_dotenv
from docx import Document
import re  # Import the re module

class ReviewSnippets:
    def __init__(self):
        load_dotenv()
        self.client = Groq(api_key=os.getenv("GROQ_API_KEY"))

    def get_blog(self, topic, structure):
        chat_completion = self.client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": f'''Generate a blog for my e-commerce site BestViewsReviews for Amazon Products. It is a GenAI platform that analyzes unstructured e-commerce product information like customer reviews, blogs, manufacturer written product descriptions, etc., and generates insights and product recommendations. 
                    Following is the structure of the blog:
                    {structure}
                    Topic: {topic}'''
                },
                {
                    "role": "user",
                    "content": f'Topic: {topic}'
                },
            ],
            model="llama3-70b-8192",
        )

        return chat_completion.choices[0].message.content

    def save_to_doc(self, content, filename):
        doc = Document()
        paragraph = doc.add_paragraph()
        
        # Split the content by double asterisks to identify bold text
        parts = re.split(r'(\*\*.*?\*\*)', content)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Add bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Add normal text
                paragraph.add_run(part)
        
        doc.save(filename)

if __name__ == "__main__":
    extractor = ReviewSnippets()
    topic = "Best Food Processors in India"
    structure = '''Blog Structure:
Introduction
 
Top picks based on important aspects – two products each aspect (include at least 4 to 5 aspects) - Provide detailed explaination of each product for each aspects 2 paragraphs for each
 
Things to keep in mind while choosing a food processor (6 to 7 factors to consider while buying a food processor)
 
Expert tips on how to maintain a food processor
 
FAQs - Provide detailed explanations for the following questions:

Can a food processor replace a blender?
Can I use a food processor for hot ingredients?
What are 3 things you can do with a food processor?
Which is better food processor or juicer?
What size of food processor do I need?
'''

    blog_content = extractor.get_blog(topic, structure=structure) 
    print(blog_content)

    # Save the blog content to a .doc file with bold formatting
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)

    filename = os.path.join(output_dir, "Blog_Content.docx")
    extractor.save_to_doc(blog_content, filename)
    print(f"Blog content saved to {filename}")
