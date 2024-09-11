import os
import requests
from dotenv import load_dotenv
from docx import Document
import re
import json
import streamlit as st
from io import BytesIO
from zipfile import ZipFile

class ReviewSnippets:
    def __init__(self):
        load_dotenv()
        self.api_url = os.getenv("Llama_API_URL")
        self.api_key = os.getenv("Llama_API_KEY")

    def get_blog(self, topic, structure):
        messages = [
            {
                "role": "system",
                "content": f'''Act as an expert content writer and generate a blog for my e-commerce site BestViewsReviews for Amazon Products. It is a GenAI platform that analyzes unstructured e-commerce product information like customer reviews, blogs, manufacturer-written product descriptions, etc., and generates insights and product recommendations. 
                Following is the structure of the blog:
                {structure}
                Topic: {topic}'''
            },
            {
                "role": "user",
                "content": f'Topic: {topic}'
            }
        ]

        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {self.api_key}'
        }
        payload = {
            "model": "meta-llama/Meta-Llama-3-8B-Instruct",
            "temperature": 0.7,
            "top_p": 0.95,
            "messages": messages,
            "max_tokens": 2048
        }

        try:
            response = requests.post(self.api_url, headers=headers, data=json.dumps(payload))
            response.raise_for_status()
            response_json = response.json()
            generated_text = response_json['choices'][0]['message']['content']
            return generated_text

        except requests.exceptions.RequestException as e:
            return f"Request failed: {e}"

    def save_to_doc(self, content, topic):
        doc = Document()
        paragraph = doc.add_paragraph()

        parts = re.split(r'(\*\*.*?\*\*)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

# Streamlit UI
def main():
    st.title("AI Blog Generator for BVR")

    mode = st.radio("Select Mode", ["Single Blog", "Multiple Blogs with single blog structure", "Multiple Blogs with separate structure"])

    if mode == "Single Blog":
        topic = st.text_input("Enter the blog topic", "Best Food Processors in India")
        structure = st.text_area("Enter the blog structure",'''Blog Structure:
        Introduction

        Top picks based on important aspects – two products each aspect (include at least 4 to 5 aspects) - Provide detailed explanation of each product for each aspect, 2 paragraphs for each

        Things to keep in mind while choosing a food processor (6 to 7 factors to consider while buying a food processor)

        Expert tips on how to maintain a food processor

        FAQs - Provide detailed explanations for the following questions:

        Can a food processor replace a blender?
        Can I use a food processor for hot ingredients?
        What are 3 things you can do with a food processor?
        Which is better food processor or juicer?
        What size of food processor do I need?
        ''')

        if st.button("Generate Blog"):
            extractor = ReviewSnippets()
            blog_content = extractor.get_blog(topic, structure)
            st.session_state.single_blog_content = blog_content

        if 'single_blog_content' in st.session_state:
            st.text_area("Generated Blog Content", st.session_state.single_blog_content, height=400)
            doc_io = ReviewSnippets().save_to_doc(st.session_state.single_blog_content, topic)
            st.download_button(
                label="Download Blog as DOCX",
                data=doc_io,
                file_name=f"{topic}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    elif mode == "Multiple Blogs with single blog structure":
        topics = st.text_area("Enter the blog topics (one per line)", "Topic 1\nTopic 2\nTopic 3")
        structure = st.text_area("Enter the blog structure", '''Blog Structure...''')

        if st.button("Generate Blogs"):
            extractor = ReviewSnippets()
            topics_list = topics.split("\n")
            st.session_state.blogs = []

            for topic in topics_list:
                blog_content = extractor.get_blog(topic, structure)
                st.session_state.blogs.append((topic, blog_content))

        if 'blogs' in st.session_state:
            zip_buffer = BytesIO()
            with ZipFile(zip_buffer, "w") as zip_file:
                for topic, blog_content in st.session_state.blogs:
                    st.text_area(f"Blog for {topic}", blog_content, height=200)
                    doc_io = ReviewSnippets().save_to_doc(blog_content, topic)
                    zip_file.writestr(f"{topic}.docx", doc_io.read())

                    st.download_button(
                        label=f"Download {topic}.docx",
                        data=doc_io,
                        file_name=f"{topic}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            zip_buffer.seek(0)
            st.download_button(
                label="Download All Blogs as ZIP",
                data=zip_buffer,
                file_name="all_blogs.zip",
                mime="application/zip"
            )

    elif mode == "Multiple Blogs with separate structure":
        num_blogs = st.number_input("Enter number of blogs", min_value=1, step=1)

        topics = []
        structures = []
        
        for i in range(num_blogs):
            topic = st.text_input(f"Enter topic for Blog {i+1}", f"Topic {i+1}")
            structure = st.text_area(f"Enter structure for Blog {i+1}", f"Structure {i+1}")
            topics.append(topic)
            structures.append(structure)

        if st.button("Generate Blogs"):
            extractor = ReviewSnippets()
            st.session_state.blogs = []

            for topic, structure in zip(topics, structures):
                blog_content = extractor.get_blog(topic, structure)
                st.session_state.blogs.append((topic, blog_content))

        if 'blogs' in st.session_state:
            zip_buffer = BytesIO()
            with ZipFile(zip_buffer, "w") as zip_file:
                for topic, blog_content in st.session_state.blogs:
                    st.text_area(f"Blog for {topic}", blog_content, height=200)
                    doc_io = ReviewSnippets().save_to_doc(blog_content, topic)
                    zip_file.writestr(f"{topic}.docx", doc_io.read())

                    st.download_button(
                        label=f"Download {topic}.docx",
                        data=doc_io,
                        file_name=f"{topic}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            zip_buffer.seek(0)
            st.download_button(
                label="Download All Blogs as ZIP",
                data=zip_buffer,
                file_name="all_blogs.zip",
                mime="application/zip"
            )

    st.markdown("---")
    st.markdown("### Created by AkshayTriapthiShorthillsAI")

if __name__ == "__main__":
    main()
